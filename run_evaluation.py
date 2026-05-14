#!/usr/bin/env python3
"""
run_evaluation.py
==================
Single-command pipeline orchestrator for SE-Agentic-Evaluator.

Runs the full evaluation pipeline in order:
  1. DOCX → Markdown conversion (docx_extract_skill)
  2. Rubric import to YAML config (evaluator_skill)
  3. Full extraction + analysis + criterion evaluation (analyze_skill)
  4. Weighted grade calculation (grade_skill)
  5. Final report generation (report_skill)

Usage:
    python run_evaluation.py \
        --input tests/test2/A1.2\ Memoria\ trabajo\ final.docx \
        --rubric tests/test2/rubrica-hito-1.md \
        --output tests/test2/output/
"""

from __future__ import annotations

import argparse
import json
import logging
import sys
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

REPO_ROOT = Path(__file__).parent


def phase0_docx_to_markdown(docx_path: Path, output_dir: Path) -> Path:
    """Phase 0: Convert DOCX to Markdown with image extraction."""
    try:
        import docx
    except ImportError:
        raise SystemExit("python-docx is required. Install with: pip install python-docx")

    md_output = output_dir / "phase0_extract" / "contents.md"
    img_dir = output_dir / "phase0_extract" / "img"
    md_output.parent.mkdir(parents=True, exist_ok=True)
    img_dir.mkdir(parents=True, exist_ok=True)

    doc = docx.Document(str(docx_path))
    lines = []
    img_count = 0

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        if "Title" in style:
            lines.append(f"# {text}")
        elif "Heading 1" in style:
            lines.append(f"## {text}")
        elif "Heading 2" in style:
            lines.append(f"### {text}")
        elif "Heading 3" in style:
            lines.append(f"#### {text}")
        elif "List" in style:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    for i, table in enumerate(doc.tables):
        lines.append(f"\n### Tabla {i+1}\n")
        header = [cell.text for cell in table.rows[0].cells]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        for row in table.rows[1:]:
            cells = [cell.text.replace("|", "\\|") for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")

    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.reltype:
            ext = rel.target_ref.split(".")[-1] if "." in rel.target_ref else "png"
            img_name = f"img_{i}.{ext}"
            img_path = img_dir / img_name
            try:
                with open(img_path, "wb") as f:
                    f.write(rel.target_part.blob)
                lines.append(f'\n![Imagen {i+1}](img/{img_name})\n')
                img_count += 1
            except Exception:
                pass

    md_content = "\n\n".join(lines)
    md_output.write_text(md_content, encoding="utf-8")
    logger.info("Phase 0: Extracted %d chars, %d images", len(md_content), img_count)
    return md_output


def phase1_import_rubric(rubric_path: Path, output_dir: Path) -> Path:
    """Phase 1: Import rubric Markdown table to YAML config."""
    sys.path.insert(0, str(REPO_ROOT))
    from core.config.rubric_importer import import_rubric

    rubric_id = rubric_path.stem.replace("rubrica-", "").replace("rubric-", "")
    config_path = REPO_ROOT / "configs" / f"rubric_{rubric_id}.yaml"
    prompts_dir = REPO_ROOT / "prompts" / rubric_id

    import_rubric(str(rubric_path), str(config_path), str(prompts_dir))

    # Fix prompt paths to include subdirectory prefix
    import yaml
    with open(config_path) as f:
        cfg = yaml.safe_load(f)
    for criterion in cfg["rubric"]["criteria"]:
        criterion["prompt"] = f"{rubric_id}/" + criterion["prompt"]
    with open(config_path, "w") as f:
        yaml.dump(cfg, f, default_flow_style=False, allow_unicode=True)

    logger.info("Phase 1: Rubric config saved to %s", config_path)
    return config_path


def phase2_evaluate(document_md: Path, config_path: Path, output_dir: Path, full: bool = False) -> dict:
    """Phase 2: Evaluate all criteria with LLM."""
    sys.path.insert(0, str(REPO_ROOT))
    from core.evaluation.criterion_evaluator import run_criterion_evaluation

    eval_dir = output_dir / "evaluacion"
    return run_criterion_evaluation(
        document_path=str(document_md),
        config_path=str(config_path),
        output_dir=str(eval_dir),
        full=full,
    )


def phase3_grade(eval_dir: Path, config_path: Path) -> dict:
    """Phase 3: Calculate weighted final grade."""
    sys.path.insert(0, str(REPO_ROOT))
    from core.grading.grader import RubricGrader

    scores_file = eval_dir / "scores.json"
    with open(scores_file) as f:
        scores_data = json.load(f)
    scores = scores_data.get("scores", {})

    grader = RubricGrader.from_config(str(config_path))
    result = grader.grade({}, scores=scores)
    return result.as_dict()


def phase4_report(document_md: Path, eval_dir: Path, config_path: Path, scores_file: Path) -> Path:
    """Phase 4: Generate final consolidated report."""
    sys.path.insert(0, str(REPO_ROOT))
    from core.evaluation.evaluator import run_report_generation

    run_report_generation(
        document_path=str(document_md),
        eval_dir=str(eval_dir),
        config_path=str(config_path),
        scores_path=str(scores_file),
        output_dir=str(eval_dir),
    )
    return eval_dir / "evaluacion_final.md"


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Run the full SE-Agentic-Evaluator pipeline."
    )
    parser.add_argument("--input", type=str, required=True, help="Path to DOCX document")
    parser.add_argument("--rubric", type=str, required=True, help="Path to rubric Markdown table")
    parser.add_argument("--output", type=str, required=True, help="Output directory")
    parser.add_argument(
        "--full", action="store_true",
        help="Run full extraction + analysis pipeline (objectives, requirements, orphans, SMART, ISO 25010)"
    )
    args = parser.parse_args()

    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    input_path = Path(args.input)
    rubric_path = Path(args.rubric)
    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)

    if not input_path.exists():
        raise SystemExit(f"Input document not found: {input_path}")
    if not rubric_path.exists():
        raise SystemExit(f"Rubric file not found: {rubric_path}")

    print(f"\n{'='*60}")
    print(f"SE-Agentic-Evaluator v2.0 — Full Pipeline")
    print(f"Documento: {input_path.name}")
    print(f"Rúbrica: {rubric_path.name}")
    print(f"Output: {output_dir}")
    print(f"Full analysis: {'Yes' if args.full else 'No'}")
    print(f"{'='*60}\n")

    # Phase 0: DOCX → Markdown
    print("[Phase 0] Converting DOCX to Markdown...")
    document_md = phase0_docx_to_markdown(input_path, output_dir)
    print(f"  → {document_md}")

    # Phase 1: Import rubric
    print("[Phase 1] Importing rubric to YAML config...")
    config_path = phase1_import_rubric(rubric_path, output_dir)
    print(f"  → {config_path}")

    # Phase 2: Evaluate criteria
    print("[Phase 2] Evaluating criteria with LLM...")
    scores_data = phase2_evaluate(document_md, config_path, output_dir, full=args.full)
    eval_dir = output_dir / "evaluacion"
    print(f"  → {eval_dir}/")

    # Phase 3: Calculate grade
    print("[Phase 3] Calculating weighted final grade...")
    grade_result = phase3_grade(eval_dir, config_path)
    print(f"  → Nota ponderada: {grade_result['weighted_final']}/10")
    print(f"  → Media (x̄): {grade_result['mean_xbar']}/10")
    print(f"  → Nivel: {grade_result['performance_level']}")

    # Phase 4: Generate report
    print("[Phase 4] Generating final report...")
    report_path = phase4_report(document_md, eval_dir, config_path, eval_dir / "scores.json")
    print(f"  → {report_path}")

    print(f"\n{'='*60}")
    print(f"Pipeline completed successfully!")
    print(f"Final grade: {grade_result['weighted_final']}/10 ({grade_result['performance_level']})")
    print(f"Report: {report_path}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
